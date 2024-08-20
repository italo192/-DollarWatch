from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import NoSuchElementException
from datetime import datetime, timedelta
from time import sleep
from docx import Document
from docx2pdf import convert
import os
import schedule
import threading

# Variável global para armazenar a última execução
ultima_execucao = None

class CotacaoApp:
    def __init__(self):
        self.log_text = ""

    def log(self, message):
        self.log_text += message + "\n\n"
        print(message)

    def iniciar_driver(self):
        self.log("Iniciando o driver do Chrome.")
        chrome_options = Options()
        chrome_options.add_argument("--window-size=800,600")
        chrome_options.add_argument("--log-level=3")
        chrome_options.add_experimental_option("excludeSwitches", ["enable-logging"])
        driver = webdriver.Chrome(options=chrome_options)
        self.log("Driver do Chrome iniciado com sucesso.")
        return driver

    def acessar_site(self, driver, url):
        self.log(f"Acessando o site: {url}")
        driver.get(url)
        sleep(5)

    def coletar_cotacao(self, driver):
        try:
            self.log("Coletando a cotação do dólar...")
            preco_element = driver.find_element(By.CSS_SELECTOR, 'div[class="text-5xl/9 font-bold text-[#232526] md:text-[42px] md:leading-[60px]"]')
            preco = preco_element.text
            data = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.log(f"Valor do Dólar: R${preco} | Data: {data}")
            return preco, data
        except NoSuchElementException:
            self.log("Elemento do preço não encontrado.")
            return None, None

    def tirar_screenshot(self, driver, path):
        self.log("Capturando a tela do site.")
        driver.save_screenshot(path)
        self.log(f"Screenshot salva em {path}")

    def criar_relatorio(self, preco, data, screenshot_path, url):
        self.log("Criando relatório em Word.")
        doc = Document()
        doc.add_heading('Relatório de Cotação do Dólar', 0)
        
        doc.add_heading('Dados da Cotação', level=1)
        doc.add_paragraph(f'Valor do Dólar: R${preco}')
        doc.add_paragraph(f'Data da Cotação: {data}')
        doc.add_paragraph(f'Site da Cotação: {url}')
        
        doc.add_heading('Screenshot do Site', level=1)
        doc.add_picture(screenshot_path)
        
        word_path = 'relatorio_cotacao.docx'
        doc.save(word_path)
        self.log(f"Relatório Word criado em {word_path}")
        return word_path

    def converter_para_pdf(self, word_path):
        self.log("Convertendo o relatório para PDF, o processo pode deromar um pouco por favor aguarde.")
        convert(word_path)
        pdf_path = word_path.replace('.docx', '.pdf')
        self.log(f"Relatório PDF criado em {pdf_path}")
        return pdf_path

    def verificar_arquivos_existentes(self):
        self.log("Verificando se já existem arquivos anteriores.")
        word_exists = os.path.exists('relatorio_cotacao.docx')
        pdf_exists = os.path.exists('relatorio_cotacao.pdf')
        return word_exists or pdf_exists

    def verificar_ultima_execucao(self):
        global ultima_execucao
        self.log("Verificando a última execução.")
        return ultima_execucao

    def salvar_ultima_execucao(self):
        global ultima_execucao
        ultima_execucao = datetime.now().date()
        self.log(f"Agendando a próxima execução. Data atual: {ultima_execucao}")

    def main(self):
        self.log("Iniciando o processo principal...")
        ultima_execucao = self.verificar_ultima_execucao()
        hoje = datetime.now().date()
        
        if ultima_execucao == hoje:
            self.log("Cotação já foi atualizada hoje.")
            return
        
        if self.verificar_arquivos_existentes():
            self.log("Arquivos existentes encontrados. Atualizando...")
        
        url = "https://www.investing.com/currencies/usd-brl"
        driver = self.iniciar_driver()
        self.acessar_site(driver, url)
        
        preco, data = self.coletar_cotacao(driver)
        
        if preco and data:
            screenshot_path = 'screenshot.png'
            self.tirar_screenshot(driver, screenshot_path)
            self.log("Cotação capturada com sucesso. Encerrando o driver.")
            driver.quit()
            
            word_path = self.criar_relatorio(preco, data, screenshot_path, url)
            pdf_path = self.converter_para_pdf(word_path)
            
            self.log(f'Relatório gerado com sucesso: {pdf_path}')
            self.salvar_ultima_execucao()
        else:
            self.log('Falha ao coletar a cotação.')
            driver.quit()

    def run_schedule(self):
        while True:
            schedule.run_pending()
            sleep(1)

if __name__ == "__main__":
    print("Pressione qualquer tecla para iniciar o programa...")
    input()
    
    app = CotacaoApp()
    
    # Agendar a execução diária
    schedule.every().day.at("09:00").do(app.main)  # Ajuste o horário conforme necessário
    
    # Executar o agendamento em uma thread separada
    threading.Thread(target=app.run_schedule).start()
    
    # Executa a coleta manualmente (se necessário)
    app.main()
